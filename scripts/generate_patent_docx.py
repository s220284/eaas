"""
Generate a formatted DOCX of the Provisional Patent Application with embedded figures.
Upload the output to Google Drive and open as Google Doc.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

DOCS_DIR = Path(__file__).parent.parent / "docs"
MD_FILE = DOCS_DIR / "PROVISIONAL_PATENT_APPLICATION.md"
OUTPUT_FILE = DOCS_DIR / "Provisional_Patent_Application.docx"

FIGURES = {
    "FIG. 1": DOCS_DIR / "Provisional_Patent_Application_Fig1.png",
    "FIG. 2": DOCS_DIR / "Provisional_Patent_Application_Fig2.png",
    "FIG. 3": DOCS_DIR / "Provisional_Patent_Application_Fig3.png",
    "FIG. 4": DOCS_DIR / "Provisional_Patent_Application_Fig4.png",
    "FIG. 5": DOCS_DIR / "Provisional_Patent_Application_Fig5.png",
    "FIG. 6": DOCS_DIR / "Provisional_Patent_Application_Fig6.png",
    "FIG. 7": DOCS_DIR / "Provisional_Patent_Application_Fig7.png",
}


def setup_styles(doc: Document):
    """Configure document styles for patent application formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 5):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = "Times New Roman"
            h.font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                h.font.size = Pt(16)
                h.font.bold = True
            elif level == 2:
                h.font.size = Pt(14)
                h.font.bold = True
            elif level == 3:
                h.font.size = Pt(13)
                h.font.bold = True
            elif level == 4:
                h.font.size = Pt(12)
                h.font.bold = True


def add_figure(doc: Document, fig_label: str, width_inches: float = 5.5):
    """Insert a figure image centered with a caption."""
    fig_path = FIGURES.get(fig_label)
    if not fig_path or not fig_path.exists():
        p = doc.add_paragraph(f"[{fig_label} — image file not found]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width_inches))

    caption = doc.add_paragraph(fig_label)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].bold = True
    caption.runs[0].font.size = Pt(10)
    doc.add_paragraph()  # spacing


def parse_and_build(doc: Document, md_text: str):
    """Parse the markdown and build the DOCX, inserting figures at appropriate points."""
    lines = md_text.split("\n")
    i = 0
    # Track which figures have been inserted so we insert each once
    inserted_figures = set()
    # Map of where figures should be inserted (after certain text patterns)
    figure_insertion_points = {
        "All services read from and write to the Data Layer (160).": "FIG. 1",
        "canon_relationships (230)": None,  # FIG. 2 handled specially
        "safety_age_gating (286)": "FIG. 2",
        "the evaluation pipeline is shown as a flow diagram:": "FIG. 3",
        "Referring now to FIG. 4, the scoring computation": "FIG. 4",
        "the entity-relationship diagram shows:": None,  # FIG. 5 after the ER section
    }

    # Skip the filing instructions section marker
    skip_filing_instructions = False

    while i < len(lines):
        line = lines[i]

        # Skip the FILING INSTRUCTIONS section entirely from the doc body —
        # we'll add it as an appendix
        if line.strip() == "## FILING INSTRUCTIONS":
            skip_filing_instructions = True
            # Add a page break before filing instructions
            doc.add_page_break()
            p = doc.add_heading("FILING INSTRUCTIONS", level=1)
            i += 1
            continue

        # Title line
        if line.startswith("# PROVISIONAL PATENT APPLICATION"):
            p = doc.add_heading("PROVISIONAL PATENT APPLICATION", level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H2 headings
        if line.startswith("## "):
            text = line[3:].strip()
            if text.startswith("Filed Under"):
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue

        # H3 headings
        if line.startswith("### "):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # H4 headings
        if line.startswith("#### "):
            text = line[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # H5 headings
        if line.startswith("##### "):
            text = line[6:].strip()
            doc.add_heading(text, level=4)
            i += 1
            continue

        # Horizontal rules — skip
        if line.strip() == "---":
            i += 1
            continue

        # Empty lines
        if line.strip() == "":
            i += 1
            continue

        # Bold key-value lines (like **Inventor(s):** ...)
        bold_match = re.match(r"^\*\*(.+?):\*\*\s*(.*)", line)
        if bold_match:
            p = doc.add_paragraph()
            run = p.add_run(bold_match.group(1) + ": ")
            run.bold = True
            p.add_run(bold_match.group(2))
            i += 1
            continue

        # Indented formula lines
        if line.strip().startswith("total_score ="):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            run.italic = True
            run.font.name = "Courier New"
            run.font.size = Pt(11)
            i += 1
            continue

        # Numbered list items (like "1. **Canon Drift**:...")
        num_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if num_match:
            text = num_match.group(2)
            text = format_inline(text)
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, num_match.group(2))
            i += 1
            continue

        # Bullet items
        if line.startswith("- ") or line.startswith("  - "):
            indent_level = 0
            if line.startswith("  - "):
                indent_level = 1
            text = line.lstrip(" -").strip()
            p = doc.add_paragraph(style="List Bullet")
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_runs(p, text)
            i += 1
            continue

        # Regular paragraph
        # Collect continuation lines
        para_text = line
        while (i + 1 < len(lines)
               and lines[i + 1].strip() != ""
               and not lines[i + 1].startswith("#")
               and not lines[i + 1].startswith("- ")
               and not lines[i + 1].startswith("**")
               and not re.match(r"^\d+\.\s+", lines[i + 1])
               and lines[i + 1].strip() != "---"
               and not lines[i + 1].strip().startswith("total_score")):
            i += 1
            para_text += " " + lines[i].strip()

        p = doc.add_paragraph()
        add_formatted_runs(p, para_text.strip())

        # Check if we should insert a figure after this paragraph
        para_plain = para_text.strip()
        for trigger, fig in figure_insertion_points.items():
            if trigger in para_plain and fig and fig not in inserted_figures:
                add_figure(doc, fig)
                inserted_figures.add(fig)

        # Insert FIG. 4 right after the paragraph that introduces it
        if "Referring now to FIG. 4" in para_plain and "FIG. 4" not in inserted_figures:
            add_figure(doc, "FIG. 4")
            inserted_figures.add("FIG. 4")

        i += 1

    # Insert any figures that weren't placed inline at the end as a drawings section
    remaining = [f for f in FIGURES if f not in inserted_figures]
    if remaining:
        doc.add_page_break()
        doc.add_heading("DRAWINGS", level=1)
        for fig_label in sorted(remaining):
            add_figure(doc, fig_label)


def format_inline(text: str) -> str:
    """Strip markdown formatting for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def add_formatted_runs(paragraph, text: str):
    """Add runs to a paragraph with bold/italic/code markdown formatting."""
    # Pattern to match **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last_end = 0
    for m in pattern.finditer(text):
        # Add plain text before match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # italic
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # code
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        last_end = m.end()

    # Add remaining plain text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def main():
    md_text = MD_FILE.read_text()

    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    setup_styles(doc)
    parse_and_build(doc, md_text)

    doc.save(str(OUTPUT_FILE))
    print(f"Created: {OUTPUT_FILE}")
    print(f"File size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")
    print()
    print("To upload to Google Docs:")
    print("1. Go to drive.google.com")
    print("2. Drag and drop the .docx file")
    print("3. Double-click to open → it will convert to Google Docs format")


if __name__ == "__main__":
    main()
